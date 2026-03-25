"""
开造 VibeBuild — 简历文件解析工具

支持格式：
- PDF (.pdf)     → pdfplumber 提取
- Word (.docx)   → python-docx 提取
- Markdown (.md) → 原文读取
- 纯文本 (.txt)  → 原文读取
"""

import io
from pathlib import Path

import structlog

logger = structlog.get_logger()

# 支持的文件 MIME 类型映射
SUPPORTED_CONTENT_TYPES = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "text/markdown": "md",
    "text/plain": "txt",
    # 常见浏览器上传变体
    "application/octet-stream": None,  # 需要通过后缀名判断
}

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".md", ".txt", ".markdown"}


def detect_file_type(filename: str, content_type: str) -> str:
    """根据文件名和 MIME 类型推断文件格式"""
    # 优先通过后缀名判断
    ext = Path(filename).suffix.lower() if filename else ""
    if ext in (".pdf",):
        return "pdf"
    if ext in (".docx",):
        return "docx"
    if ext in (".md", ".markdown"):
        return "md"
    if ext in (".txt",):
        return "txt"

    # 通过 MIME 类型判断
    mapped = SUPPORTED_CONTENT_TYPES.get(content_type)
    if mapped:
        return mapped

    return "txt"  # 默认当纯文本处理


async def parse_resume_file(file_bytes: bytes, filename: str, content_type: str) -> str:
    """
    解析简历文件，提取纯文本内容

    Args:
        file_bytes: 文件二进制内容
        filename: 原始文件名
        content_type: MIME 类型

    Returns:
        提取出的纯文本字符串

    Raises:
        ValueError: 文件格式不支持或解析失败
    """
    file_type = detect_file_type(filename, content_type)
    logger.info("resume_parse_start", filename=filename, content_type=content_type, detected_type=file_type)

    try:
        if file_type == "pdf":
            return _parse_pdf(file_bytes)
        elif file_type == "docx":
            return _parse_docx(file_bytes)
        elif file_type in ("md", "txt"):
            return _parse_text(file_bytes)
        else:
            raise ValueError(f"不支持的文件格式: {filename} ({content_type})")
    except ValueError:
        raise
    except Exception as e:
        logger.error("resume_parse_error", filename=filename, error=str(e))
        raise ValueError(f"文件解析失败: {e}")


def _parse_pdf(file_bytes: bytes) -> str:
    """解析 PDF 文件"""
    import pdfplumber

    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text.strip())

    result = "\n\n".join(text_parts)
    if not result.strip():
        raise ValueError("PDF 文件未提取到有效文本内容（可能是扫描件或纯图片 PDF）")
    return result


def _parse_docx(file_bytes: bytes) -> str:
    """解析 Word (.docx) 文件"""
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    text_parts = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            text_parts.append(text)

    # 也提取表格内容（简历常用表格排版）
    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                text_parts.append(" | ".join(row_texts))

    result = "\n".join(text_parts)
    if not result.strip():
        raise ValueError("Word 文件未提取到有效文本内容")
    return result


def _parse_text(file_bytes: bytes) -> str:
    """解析纯文本 / Markdown 文件"""
    # 尝试 UTF-8，退回 GBK（中文简历常见编码）
    for encoding in ("utf-8", "gbk", "gb2312", "latin-1"):
        try:
            return file_bytes.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            continue
    raise ValueError("文件编码无法识别，请使用 UTF-8 编码保存")
